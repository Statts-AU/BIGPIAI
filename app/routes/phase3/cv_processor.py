# app/routes/phase3/cv_processor.py
"""
Optimized CV processing logic using simplified pipeline
Based on cv_processor reference with 1 API call per CV
"""
import sys
import time
from pathlib import Path
from typing import List, Tuple, Dict, Any
from flask import current_app

# Add project root to Python path
project_root = Path(__file__).parent.parent.parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

try:
    from app.cv_processor.analysis.placeholder_mapper import (
        extract_jinja_placeholders,
        map_jinja_placeholders_to_values
    )
    from app.cv_processor.analysis.cv_parser import extract_raw_cv_text, _read_docx_full_text, convert_pdf_to_docx
    from app.cv_processor.analysis.template_analyzer import read_template_full_text
    CV_PROCESSOR_AVAILABLE = True
    print("✓ CV Processor modules loaded successfully")
except ImportError as e:
    print(f"⚠ CV Processor not available: {e}")
    CV_PROCESSOR_AVAILABLE = False

from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

def extract_template_placeholders(template_path: Path) -> Tuple[str, List[str]]:
    """
    Step 1: Extract Jinja placeholders from template
    Returns: (template_text, jinja_placeholders)
    """
    if CV_PROCESSOR_AVAILABLE:
        print("[DEBUG] Using CV Processor for template analysis")
        template_text = read_template_full_text(str(template_path))
        jinja_placeholders = extract_jinja_placeholders(template_text)
        print(f"[DEBUG] Found {len(jinja_placeholders)} Jinja placeholders: {jinja_placeholders}")
    else:
        print("[DEBUG] Using fallback for template analysis")
        template_text = read_template_fallback(template_path)
        jinja_placeholders = extract_placeholders_fallback(template_text)
        print(f"[DEBUG] Found {len(jinja_placeholders)} placeholders (fallback): {jinja_placeholders}")
    
    return template_text, jinja_placeholders

def extract_cv_texts(cv_paths: List[Path]) -> List[Tuple[str, str]]:
    """
    Step 2: Extract text from CV files
    Returns: List of (filepath, resume_text) tuples
    """
    if CV_PROCESSOR_AVAILABLE:
        print("[DEBUG] Using CV Processor for text extraction")
        return extract_cv_text_with_processor(cv_paths)
    else:
        print("[DEBUG] Using fallback for text extraction")
        return extract_cv_text_fallback(cv_paths)

def map_cv_to_template(jinja_placeholders: List[str], resume_text: str, template_text: str) -> Dict[str, str]:
    """
    Step 3: Map CV data to template placeholders using AI (1 API call per CV)
    Uses the optimized Gemini prompt from cv_processor
    """
    if CV_PROCESSOR_AVAILABLE:
        print("[DEBUG] Making AI API call to map placeholders...")
        # This uses the optimized prompt from placeholder_mapper.py
        placeholder_to_value = map_jinja_placeholders_to_values(
            jinja_placeholders, resume_text, template_text
        )
        print(f"[DEBUG] AI mapping result: {list(placeholder_to_value.keys())}")
        return placeholder_to_value
    else:
        print("[DEBUG] Using fallback placeholder mapping")
        return {placeholder: f"Sample {placeholder}" for placeholder in jinja_placeholders}

def generate_document(template_path: Path, placeholder_values: Dict[str, str], output_path: Path) -> None:
    """
    Step 4: Generate document using DocxTemplate
    """
    print(f"[DEBUG] Generating document: {output_path.name}")
    
    # Generate document
    doc = DocxTemplate(str(template_path))
    doc.render(placeholder_values)
    
    # Apply post-processing for better formatting
    format_document(doc)
    
    doc.save(str(output_path))
    print(f"[DEBUG] Successfully generated: {output_path.name}")

def format_document(doc):
    """Apply formatting improvements to the document"""
    try:
        paragraphs_to_remove = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Justify paragraphs for full width
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Prevent page breaks for headings
                if 'heading' in paragraph.style.name.lower():
                    paragraph.paragraph_format.keep_with_next = True
                
                # Convert text lists to DOCX bullet lists
                if '•' in paragraph.text or any(line.strip().startswith('- ') for line in paragraph.text.split('\n')):
                    lines = paragraph.text.split('\n')
                    list_items = []
                    for line in lines:
                        line = line.strip()
                        if line.startswith('•'):
                            list_items.append(line[1:].strip())
                        elif line.startswith('- '):
                            list_items.append(line[2:].strip())
                        elif line and not any(line.startswith(marker) for marker in ['•', '- ']):
                            list_items = []
                            break
                    
                    if list_items:
                        paragraph.clear()
                        for item in list_items:
                            if item:
                                p = paragraph.insert_paragraph_before(item)
                                p.style = 'List Bullet'
                        if not paragraph.text.strip():
                            paragraphs_to_remove.append(paragraph)
        
        # Remove empty paragraphs
        for p in paragraphs_to_remove:
            if p in doc.paragraphs:
                p._element.getparent().remove(p._element)
        
        # Prevent tables from splitting across pages
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            paragraph.paragraph_format.keep_with_next = True
                            paragraph.paragraph_format.keep_together = True
                            
    except Exception as e:
        current_app.logger.warning(f"Document formatting warning: {str(e)}")

# Fallback functions when cv_processor is not available
def read_template_fallback(template_path: Path) -> str:
    """Fallback template reader using python-docx"""
    try:
        from docx import Document
        doc = Document(template_path)
        parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            parts.append(paragraph.text)
        
        return '\n'.join(parts)
    except Exception as e:
        print(f"Error reading template: {e}")
        return f"Sample template content from {template_path.name}"

def extract_placeholders_fallback(text: str) -> List[str]:
    """Fallback placeholder extraction using regex"""
    import re
    # More comprehensive pattern to catch Jinja placeholders
    pattern = r'\{\{\s*([^}]+?)\s*\}\}'
    matches = re.findall(pattern, text)
    placeholders = []
    for match in matches:
        # Handle filters and clean up
        placeholder = match.split('|')[0].strip()
        if placeholder and placeholder not in placeholders:
            placeholders.append(placeholder)
    
    # Also look for common placeholder patterns like [name], {name}, etc.
    if not placeholders:
        bracket_pattern = r'\[([^\]]+)\]'
        bracket_matches = re.findall(bracket_pattern, text)
        for match in bracket_matches:
            if match.strip() and match.strip() not in placeholders:
                placeholders.append(match.strip())
    
    return placeholders

def extract_cv_text_with_processor(cv_paths: List[Path]) -> List[Tuple[str, str]]:
    """Extract CV text using the cv_processor module"""
    results = []
    for path in cv_paths:
        try:
            if path.suffix.lower() == '.docx':
                text = _read_docx_full_text(str(path))
                results.append((str(path), text))
            elif path.suffix.lower() == '.pdf':
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_docx = temp_file.name
                convert_pdf_to_docx(str(path), temp_docx)
                text = _read_docx_full_text(temp_docx)
                results.append((str(path), text))
                os.unlink(temp_docx)
        except Exception as e:
            print(f"Error processing {path}: {e}")
            results.append((str(path), f"Error reading {path.name}"))
    
    return results

def extract_cv_text_fallback(cv_paths: List[Path]) -> List[Tuple[str, str]]:
    """Fallback CV text extraction"""
    results = []
    for path in cv_paths:
        try:
            if path.suffix.lower() == '.docx':
                from docx import Document
                doc = Document(path)
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                text = '\n'.join(text_parts)
                results.append((str(path), text))
            elif path.suffix.lower() == '.pdf':
                # For PDF, we'd need additional libraries
                results.append((str(path), f"PDF content from {path.name} (processing not available)"))
        except Exception as e:
            print(f"Error processing {path}: {e}")
            results.append((str(path), f"Error reading {path.name}"))
    return results
