# app/cv_processor/analysis/cv_parser.py
"""
CV text extraction and parsing utilities.
"""
from pathlib import Path
from typing import Tuple, Any, Dict, List
from docx import Document as DocxDocument
import tempfile
import os

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    Converter = None

from app.cv_processor.config.settings import get_content_dir
from app.cv_processor.utils.ai_client import analyze_with_ai, parse_with_ai


def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    """Convert PDF to DOCX."""
    if not PDF2DOCX_AVAILABLE:
        raise RuntimeError("pdf2docx package not installed. Run: pip install pdf2docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()


def _read_docx_full_text(file_path: str) -> str:
    """Read every bit of text from a DOCX: paragraphs, tables, headers, footers with enhanced extraction."""
    doc = DocxDocument(file_path)
    parts: list[str] = []

    # Enhanced paragraph extraction with run-level text
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
        # Also extract text from runs in case paragraph.text misses formatting
        run_texts = [run.text for run in p.runs if run.text and run.text.strip()]
        if run_texts and not p.text.strip():
            parts.append("".join(run_texts).strip())

    # Enhanced table extraction
    for t in doc.tables:
        for r in t.rows:
            row_cells: list[str] = []
            for c in r.cells:
                # Extract from both cell paragraphs and runs
                cell_parts = []
                for p in c.paragraphs:
                    if p.text and p.text.strip():
                        cell_parts.append(p.text.strip())
                    else:
                        run_text = "".join(run.text for run in p.runs if run.text)
                        if run_text.strip():
                            cell_parts.append(run_text.strip())
                if cell_parts:
                    row_cells.append(" ".join(cell_parts))
            if row_cells:
                parts.append("\t".join(row_cells))

    # Enhanced headers/footers extraction
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for p in section.footer.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())

    return "\n".join(parts)


def extract_raw_cv_text(content_dir: str = None) -> List[Tuple[str, str]]:
    """Extract raw text from all .docx and .pdf files in content_dir. Returns list of (filename, text)."""
    if content_dir is None:
        content_dir = get_content_dir()
    p = Path(content_dir)
    if not p.exists():
        return []
    results: List[Tuple[str, str]] = []
    for f in p.iterdir():
        if f.is_file():
            if f.suffix.lower() == ".docx":
                text = _read_docx_full_text(str(f))
                results.append((str(f), text))
            elif f.suffix.lower() == ".pdf":
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_docx = temp_file.name
                convert_pdf_to_docx(str(f), temp_docx)
                text = _read_docx_full_text(temp_docx)
                results.append((str(f), text))
                os.unlink(temp_docx)
    return results


def parse_resume_with_ai(raw_text: str) -> Dict[str, Any]:
    """Send raw CV text to AI to obtain a predictable structured JSON."""
    instructions = (
        "You are an expert resume parser. Extract ALL information from this CV text into a comprehensive JSON object.\n\n"
        "Required structure:\n"
        "- 'personal_info': {'full_name', 'email', 'phone', 'address', 'linkedin_url'}\n"
        "- 'current_position': current job title and company\n"
        "- 'qualifications': degrees, certifications, professional qualifications\n"
        "- 'summary': professional summary or objective\n"
        "- 'work_experience': list of jobs with 'job_title', 'company', 'dates', 'responsibilities', 'achievements'\n"
        "- 'education': list with 'degree', 'institution', 'dates', 'details'\n"
        "- 'skills': technical and soft skills\n"
        "- 'projects': significant projects with descriptions\n"
        "- 'certifications': professional certifications\n"
        "- 'languages': spoken languages\n"
        "- 'references': reference contacts if available\n\n"
        "Extract EVERY piece of information, even if it seems minor. Be thorough."
    )

    return parse_with_ai(raw_text, instructions, temperature=0.0)


def build_structured_cv_json(content_dir: str = None) -> Dict[str, Any]:
    """Top-level helper to extract raw CV text from content/ and parse via AI."""
    results = extract_raw_cv_text(content_dir)
    if not results:
        raise RuntimeError("No .docx or .pdf files found in content/ to parse.")
    raw = "\n\n---\n\n".join([text for _, text in results])
    if not raw.strip():
        raise RuntimeError("Extracted CV text is empty.")
    return parse_resume_with_ai(raw)


def build_source_profile(content_dir: str = None) -> Dict[str, Any]:
    """Process source CV(s) into a standardized, predictable JSON format (source_profile)."""
    results = extract_raw_cv_text(content_dir)
    if not results:
        raise RuntimeError("No .docx or .pdf files found in content/ to parse.")
    raw_text = "\n\n---\n\n".join([text for _, text in results])
    if not raw_text.strip():
        raise RuntimeError("Extracted CV text is empty.")

    system = "You are an expert resume parser. Return only valid JSON with no commentary."
    user_prompt = (
        "You are an expert resume parser. Create a comprehensive source profile from this CV text.\n\n"
        "Extract ALL available information into this JSON structure:\n"
        "- 'personal_info': {'full_name', 'email', 'phone', 'address', 'linkedin_url'}\n"
        "- 'current_position': current role and company\n"
        "- 'qualifications': all degrees, certifications, professional qualifications\n"
        "- 'professional_summary': career summary or objective\n"
        "- 'work_experience': comprehensive job history with all details\n"
        "- 'education': complete educational background\n"
        "- 'technical_skills': all technical competencies\n"
        "- 'projects': significant projects and achievements\n"
        "- 'certifications': professional certifications and licenses\n"
        "- 'languages': language proficiencies\n"
        "- 'availability': any availability information\n"
        "- 'references': reference contacts and details\n\n"
        "Be extremely thorough - extract every detail that could be relevant for any professional context.\n\n"
        f"CV TEXT:\n---\n{raw_text}\n---"
    )

    return analyze_with_ai(user_prompt, system)


def get_raw_cv_text(content_dir: str = None) -> str:
    """Get raw CV text for processing."""
    results = extract_raw_cv_text(content_dir)
    if not results:
        raise RuntimeError("No .docx or .pdf files found in content/ to parse.")
    texts = [text for _, text in results]
    return "\n\n---\n\n".join(texts)
