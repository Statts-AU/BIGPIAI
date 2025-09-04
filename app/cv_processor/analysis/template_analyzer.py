# app/cv_processor/analysis/template_analyzer.py
"""
Template analysis and schema building utilities.
"""
from typing import Any, Dict, List
from docx import Document as DocxDocument

from app.cv_processor.config.settings import get_template_path
from app.cv_processor.utils.ai_client import analyze_with_ai


def read_template_full_text(template_path: str = None) -> str:
    """Read all text from the template (paragraphs, tables, headers, footers)."""
    if template_path is None:
        template_path = get_template_path()
    doc = DocxDocument(template_path)
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            row_cells: list[str] = []
            for c in r.cells:
                txt = "\n".join(p.text for p in c.paragraphs if p.text)
                if txt:
                    row_cells.append(txt)
            if row_cells:
                parts.append("\t".join(row_cells))
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text:
                parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text:
                parts.append(p.text)
    return "\n".join(parts)


def analyze_template_schema(template_text: str = None, template_path: str = None) -> List[Dict[str, Any]]:
    """Analyze the template text and identify placeholders with context into a schema list."""
    text = template_text if template_text is not None else read_template_full_text(template_path)
    if not text.strip():
        raise RuntimeError("Template text is empty.")

    print("[AI] Starting template schema analysis (Step 2)...")
    system = "You are a document template analyst. Return only valid JSON with no commentary."
    user_prompt = (
        "You are an expert document template analyst. Your task is to analyze the provided template text and identify every distinct placeholder that requires user data.\n\n"
        "**--- CRITICAL RULES ---**\n\n"
        "**1. DEFINE A 'PLACEHOLDER':**\n"
        "A placeholder is a token or space intended for data replacement. Examples:\n"
        "- Explicit tokens: `xxx`, `[insert name]`, `____`, `...`\n"
        "- Instructional text in brackets: `[Summarise your experience here]`\n"
        "- Empty spaces in a table cell next to a label.\n\n"
        "**2. WHAT IS NOT A PLACEHOLDER (DO NOT EXTRACT):**\n"
        "- **Section Headings:** 'Work Experience', 'Professional Summary'\n"
        "- **Static Labels:** 'Name:', 'Phone:', 'Current Position:' (The placeholder is what comes *after* these labels).\n"
        "- **Instructional Sentences:** **'Project/s committed to and estimated delivery timeline'** and **'Project/s in procurement/project tender phase'** are complex labels/instructions, NOT placeholders themselves. The placeholder is the space or token where the user writes the answer.\n"
        "- **Formatted Text** used for structure.\n\n"
        "**3. TABLE-SPECIFIC RULE:**\n"
        "- When you see a two-column table structure like `Project#1 | xxx`, the text in the left column (`Project#1`) is the **nearby_label**, and the text in the right column (`xxx`) is the **placeholder_text**. Do NOT identify `Project#1` as a placeholder.\n\n"
        "**4. UNIQUENESS:**\n"
        "- Treat every single placeholder you find as a unique item, even if the text is identical (e.g., two separate `xxx` fields). Assign a sequential `unique_id` to each one.\n\n"
        "**--- OUTPUT FORMAT ---**\n"
        "For each TRUE placeholder, create a JSON object with this exact shape:\n"
        "- `unique_id`: A unique sequential number for this specific placeholder instance (e.g., 1, 2, 3...).\n"
        "- `placeholder_text`: The exact text of the placeholder to be replaced (e.g., 'xxx', '[insert role]').\n"
        "- `nearby_label`: The label text immediately preceding or adjacent to the placeholder that gives it context (e.g., 'Name:', 'Project#1').\n"
        "- `context_description`: A detailed explanation of the information needed based on the placeholder and its surrounding text.\n"
        "- `field_type`: Classify the field as one of: 'personal_info', 'qualification', 'experience', 'project', 'reference', 'availability', 'summary', 'other'.\n"
        "- `expected_content`: Describe the expected data format (e.g., 'A full name', 'A project description with timeline', 'A single date').\n\n"
        "**--- TEMPLATE TEXT TO ANALYZE ---**\n"
        f"{text}"
    )

    result = analyze_with_ai(user_prompt, system)
    if not isinstance(result, list):
        raise RuntimeError("template_schema is not a JSON list.")
    print("[AI] Template schema analysis complete.")
    return result


def build_allowed_keys_from_resume(resume: Dict[str, Any]) -> List[str]:
    """Derive a set of allowed semantic keys from the parsed resume JSON."""
    keys: list[str] = []
    base_candidates = [
        "full_name",
        "email",
        "phone",
        "linkedin_url",
        "summary",
    ]
    for k in base_candidates:
        keys.append(k)

    # Work experience
    work = resume.get("work_experience") or []
    if isinstance(work, list):
        for i, item in enumerate(work, start=1):
            if not isinstance(item, dict):
                continue
            # Prefer common keys; fall back to any present keys
            job_keys = ["job_title", "title", "position"]
            company_keys = ["company", "employer", "organization", "organisation"]
            dates_keys = ["dates", "period", "start_end"]
            resp_keys = ["responsibilities", "description", "highlights"]
            if any(k in item for k in job_keys):
                keys.append(f"job_title_{i}")
            if any(k in item for k in company_keys):
                keys.append(f"company_{i}")
            if any(k in item for k in dates_keys):
                keys.append(f"dates_{i}")
            if any(k in item for k in resp_keys):
                keys.append(f"responsibilities_{i}")

    # Education
    edu = resume.get("education") or []
    if isinstance(edu, list):
        for i, item in enumerate(edu, start=1):
            if not isinstance(item, dict):
                continue
            degree_keys = ["degree", "qualification"]
            inst_keys = ["institution", "school", "university", "college"]
            dates_keys = ["dates", "period", "year"]
            if any(k in item for k in degree_keys):
                keys.append(f"degree_{i}")
            if any(k in item for k in inst_keys):
                keys.append(f"institution_{i}")
            if any(k in item for k in dates_keys):
                keys.append(f"education_dates_{i}")

    # Deduplicate preserving order
    seen: set[str] = set()
    out: list[str] = []
    for k in keys:
        if k not in seen:
            seen.add(k)
            out.append(k)
    return out
