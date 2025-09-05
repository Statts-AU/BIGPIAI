# cv_processor/editor/python_docx_editor.py
"""
Python-DOCX based editor implementation.
"""
from typing import Dict, Any, Tuple

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    PYTHON_DOCX_AVAILABLE = False

from app.cv_processor.editor.base_editor import BaseDocxEditor


class PythonDocxEditor(BaseDocxEditor):
    """DOCX editor using python-docx package."""

    def __init__(self, file_path: str):
        super().__init__(file_path)
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx package is not available.")
        self.doc = DocxDocument(file_path)
        self.editor_type = "python-docx"

    def get_enumerated_structure(self) -> Tuple[str, Dict[str, Any]]:
        """Extract document content with unique positional identifiers using python-docx."""
        enumerated_lines = []
        position_map = {}

        # Process paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            para_id = f"PARAGRAPH_{i}"
            text = paragraph.text.strip()
            if text:
                enumerated_lines.append(f"{para_id}::{text}")
                position_map[para_id] = {
                    "type": "paragraph",
                    "index": i,
                    "text": text
                }

        # Process tables
        for t_idx, table in enumerate(self.doc.tables):
            table_id = f"TABLE_{t_idx}"
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        cell_id = f"{table_id}_ROW_{r_idx}_COL_{c_idx}"
                        enumerated_lines.append(f"{cell_id}::{cell_text}")
                        position_map[cell_id] = {
                            "type": "table_cell",
                            "table_index": t_idx,
                            "row": r_idx,
                            "col": c_idx,
                            "text": cell_text
                        }

        enumerated_text = "\n".join(enumerated_lines)
        return enumerated_text, position_map

    def replace_at_position(self, location_id: str, placeholder: str, replacement: str, position_map: Dict[str, Any]) -> bool:
        """Replace text at specific document position using python-docx."""
        if location_id not in position_map:
            return False

        position_info = position_map[location_id]
        element_type = position_info["type"]

        if element_type == "paragraph":
            para_index = position_info["index"]
            if para_index < len(self.doc.paragraphs):
                paragraph = self.doc.paragraphs[para_index]
                original_text = paragraph.text

                # Skip if this looks like a heading or bold label
                if self._is_heading_or_label(paragraph):
                    return False

                if placeholder in original_text:
                    # Preserve formatting by replacing text in runs
                    return self._replace_in_runs(paragraph.runs, placeholder, replacement)

        elif element_type == "table_cell":
            table_idx = position_info["table_index"]
            row_idx = position_info["row"]
            col_idx = position_info["col"]

            if table_idx < len(self.doc.tables):
                table = self.doc.tables[table_idx]
                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    original_text = cell.text
                    if placeholder in original_text:
                        # Preserve formatting in table cells too
                        cell_replaced = False
                        for paragraph in cell.paragraphs:
                            if self._replace_in_runs(paragraph.runs, placeholder, replacement):
                                cell_replaced = True
                        return cell_replaced

        return False

    def save(self, output_path: str = None) -> None:
        """Save the document using python-docx."""
        save_path = output_path or self.file_path
        self.doc.save(save_path)

    def get_text(self) -> str:
        """Get all text from the document using python-docx."""
        text_parts = []
        for paragraph in self.doc.paragraphs:
            text_parts.append(paragraph.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)
