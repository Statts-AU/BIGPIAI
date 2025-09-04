"""
Core CV processing functions adapted for web uploads.
"""
from pathlib import Path
from typing import Dict, Any, List, Tuple
import tempfile
import os
from docx import Document as DocxDocument

# Import with fallback for development
try:
    from docxtpl import DocxTemplate
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False
    DocxTemplate = None

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    Converter = None

try:
    from jinja2 import Environment, meta
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False
    Environment = None
    meta = None

try:
    from app.cv_processor.utils.ai_client import analyze_with_ai, parse_with_ai
    from app.cv_processor.analysis.cv_parser import extract_raw_cv_text, _read_docx_full_text
    from app.cv_processor.analysis.template_analyzer import read_template_full_text
    from app.cv_processor.analysis.placeholder_mapper import map_jinja_placeholders_to_values, extract_jinja_placeholders
    AI_CLIENT_AVAILABLE = True
except ImportError:
    # Fallback for development
    AI_CLIENT_AVAILABLE = False
    def analyze_with_ai(prompt, system=""):
        return {"error": "AI client not available"}
    
    def parse_with_ai(text, instructions, temperature=0.0):
        return {"error": "AI client not available"}
    
    def extract_raw_cv_text(content_dir=None):
        return []
    
    def _read_docx_full_text(file_path):
        return ""
    
    def read_template_full_text(template_path=None):
        return ""
    
    def map_jinja_placeholders_to_values(placeholders, resume_text, template_text):
        return {}
    
    def extract_jinja_placeholders(template_text):
        return []


def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    """Convert PDF to DOCX."""
    if not PDF2DOCX_AVAILABLE:
        raise RuntimeError("pdf2docx package not installed. Run: pip install pdf2docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()


def _read_docx_full_text(file_path: str) -> str:
    """Read every bit of text from a DOCX: paragraphs, tables, headers, footers with enhanced extraction."""
    doc = DocxDocument(file_path)
    parts: list[str] = []

    # Enhanced paragraph extraction with run-level text
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
        # Also extract text from runs in case paragraph.text misses formatting
        run_texts = [run.text for run in p.runs if run.text and run.text.strip()]
        if run_texts and not p.text.strip():
            parts.append("".join(run_texts).strip())

    # Enhanced table extraction
    for t in doc.tables:
        for r in t.rows:
            row_cells: list[str] = []
            for c in r.cells:
                # Extract from both cell paragraphs and runs
                cell_parts = []
                for p in c.paragraphs:
                    if p.text and p.text.strip():
                        cell_parts.append(p.text.strip())
                    else:
                        run_text = "".join(run.text for run in p.runs if run.text)
                        if run_text.strip():
                            cell_parts.append(run_text.strip())
                if cell_parts:
                    row_cells.append(" ".join(cell_parts))
            if row_cells:
                parts.append("\t".join(row_cells))

    # Enhanced headers/footers extraction
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for p in section.footer.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())

    return "\n".join(parts)


def extract_raw_cv_text_from_files(file_paths: List[Path]) -> List[Tuple[str, str]]:
    """Extract raw text from uploaded CV files. Returns list of (filename, text)."""
    results: List[Tuple[str, str]] = []
    
    for file_path in file_paths:
        if not file_path.exists():
            continue
            
        try:
            if file_path.suffix.lower() == ".docx":
                text = _read_docx_full_text(str(file_path))
                results.append((str(file_path), text))
            elif file_path.suffix.lower() == ".pdf":
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_docx = temp_file.name
                convert_pdf_to_docx(str(file_path), temp_docx)
                text = _read_docx_full_text(temp_docx)
                results.append((str(file_path), text))
                os.unlink(temp_docx)
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            continue
            
    return results


def read_template_full_text_from_file(template_path: Path) -> str:
    """Read all text from the template file (paragraphs, tables, headers, footers)."""
    doc = DocxDocument(str(template_path))
    parts: list[str] = []
    
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
            
    for t in doc.tables:
        for r in t.rows:
            row_cells: list[str] = []
            for c in r.cells:
                txt = "\n".join(p.text for p in c.paragraphs if p.text)
                if txt:
                    row_cells.append(txt)
            if row_cells:
                parts.append("\t".join(row_cells))
                
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text:
                parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text:
                parts.append(p.text)
                
    return "\n".join(parts)


def extract_jinja_placeholders(template_text: str) -> List[str]:
    """Extract Jinja2 placeholders from template text."""
    # Pattern to match {{ variable_name }} placeholders
    pattern = r'\{\{\s*([^}]+)\s*\}\}'
    matches = re.findall(pattern, template_text)
    
    # Clean up the matches (remove whitespace, filters, etc.)
    placeholders = []
    for match in matches:
        # Remove filters and other Jinja syntax, keep just the variable name
        clean_match = match.split('|')[0].strip()
        if clean_match and clean_match not in placeholders:
            placeholders.append(clean_match)
    
    return placeholders


def map_jinja_placeholders_to_values(placeholders: List[str], resume_text: str, template_text: str) -> Dict[str, str]:
    """Map Jinja placeholders to values extracted from resume text using AI."""
    if not placeholders:
        return {}
    
    # Create a mapping prompt for AI
    placeholders_str = "\n".join([f"- {placeholder}" for placeholder in placeholders])
    
    instructions = f"""
    You are an expert CV data extractor. Given a resume text and a list of template placeholders, 
    extract the appropriate information from the resume to fill each placeholder.
    
    Template placeholders to fill:
    {placeholders_str}
    
    Rules:
    1. Extract exact information from the resume text
    2. If information is not available, use an empty string ""
    3. For dates, use the format found in the resume
    4. For lists (like skills), create comma-separated values
    5. Keep formatting simple and clean
    6. Return a JSON object with placeholder names as keys and extracted values as values
    
    Resume text:
    {resume_text}
    """
    
    try:
        result = parse_with_ai(resume_text, instructions, temperature=0.0)
        
        # Ensure we have a dictionary and all placeholders are covered
        if not isinstance(result, dict):
            result = {}
        
        # Fill missing placeholders with empty strings
        final_mapping = {}
        for placeholder in placeholders:
            final_mapping[placeholder] = result.get(placeholder, "")
        
        return final_mapping
        
    except Exception as e:
        print(f"Error in AI mapping: {e}")
        # Fallback: return empty strings for all placeholders
        return {placeholder: "" for placeholder in placeholders}


# Legacy function names for compatibility
def extract_raw_cv_text(content_dir: str = None) -> List[Tuple[str, str]]:
    """Legacy function - not used in web version."""
    return []

def read_template_full_text(template_path: str = None) -> str:
    """Legacy function - not used in web version."""
    return ""

def parse_resume_with_ai(raw_text: str) -> Dict[str, Any]:
    """Send raw CV text to AI to obtain a predictable structured JSON."""
    instructions = (
        "You are an expert resume parser. Extract ALL information from this CV text into a comprehensive JSON object.\n\n"
        "Required structure:\n"
        "- 'personal_info': {'full_name', 'email', 'phone', 'address', 'linkedin_url'}\n"
        "- 'current_position': current job title and company\n"
        "- 'qualifications': degrees, certifications, professional qualifications\n"
        "- 'summary': professional summary or objective\n"
        "- 'work_experience': list of jobs with 'job_title', 'company', 'dates', 'responsibilities', 'achievements'\n"
        "- 'education': list with 'degree', 'institution', 'dates', 'details'\n"
        "- 'skills': technical and soft skills\n"
        "- 'projects': significant projects with descriptions\n"
        "- 'certifications': professional certifications\n"
        "- 'languages': spoken languages\n"
        "- 'references': reference contacts if available\n\n"
        "Extract EVERY piece of information, even if it seems minor. Be thorough."
    )

    return parse_with_ai(raw_text, instructions, temperature=0.0)

def build_structured_cv_json() -> Dict[str, Any]:
    """Legacy function - not used in web version."""
    return {}

def get_raw_cv_text() -> str:
    """Legacy function - not used in web version."""
    return ""
